"""
document_loader.py — Multi-format document parser.

WHAT THIS FILE DOES:
    Reads raw files (PDF, DOCX, TXT, CSV) and converts them into
    a list of LangChain Document objects with clean text and metadata.

WHY IT EXISTS:
    Every file format stores text differently:
    - PDF  : binary format with fonts, layouts, page objects
    - DOCX : XML inside a zip archive
    - TXT  : plain UTF-8 text
    - CSV  : tabular rows and columns

    Each needs a different parsing library. This module hides that
    complexity behind one clean interface: load_document(file_path).

INDUSTRY USAGE:
    Every enterprise RAG system has a loader layer. LangChain ships
    50+ loaders (Confluence, Notion, SharePoint, S3, Google Drive).
    We build our own so you understand what happens under the hood.

DATA FLOW:
    file_path (Path)
        → detect format by extension
        → call the right parser
        → return List[Document]

    Each Document:
        page_content : the raw text of one page or section
        metadata     : {source, page, file_type}
"""

from pathlib import Path
from typing import List

from langchain_core.documents import Document

from src.config import SUPPORTED_EXTENSIONS
from src.logger import get_logger

logger = get_logger(__name__)


# ---------------------------------------------------------------------------
# Public interface — the only function other modules need to call
# ---------------------------------------------------------------------------

def load_document(file_path: str | Path) -> List[Document]:
    """
    Load a document from disk and return a list of LangChain Documents.

    One Document per PAGE for PDFs (preserves page citations).
    One Document for TXT and DOCX (split later by chunker).
    One Document per ROW for CSV (each row is a discrete record).

    Args:
        file_path: Absolute or relative path to the file.

    Returns:
        List of Document objects, each with page_content and metadata.

    Raises:
        ValueError: If the file extension is not supported.
        FileNotFoundError: If the file does not exist.
    """
    file_path = Path(file_path)

    # Fail fast with clear messages instead of cryptic library errors
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = file_path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported: {SUPPORTED_EXTENSIONS}"
        )

    logger.info("Loading document: %s (type: %s)", file_path.name, ext)

    # Dispatch to the right parser based on file extension
    # This is the Strategy pattern — swap parsers without changing callers
    parsers = {
        ".pdf":  _load_pdf,
        ".docx": _load_docx,
        ".txt":  _load_txt,
        ".csv":  _load_csv,
    }

    documents = parsers[ext](file_path)

    logger.info(
        "Loaded '%s' → %d document section(s)", file_path.name, len(documents)
    )
    return documents


# ---------------------------------------------------------------------------
# Private parsers — one per format
# ---------------------------------------------------------------------------

def _load_pdf(file_path: Path) -> List[Document]:
    """
    Parse a PDF using pypdf.

    WHY pypdf: Lightweight, pure Python, handles most PDFs well.
    Alternative: pdfplumber (better for tables), pymupdf (fastest).

    IMPORTANT: We create one Document per PAGE.
    This preserves page numbers in metadata for page-level citations.
    "This answer came from page 12 of annual_report.pdf"
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf not installed. Run: pip install pypdf")

    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError:
        convert_from_path = None
        pytesseract = None

    reader = PdfReader(str(file_path))
    documents = []

    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()

        # Fallback to OCR if page seems to be a scanned image (no text)
        if len(text) < 10 and convert_from_path and pytesseract:
            logger.debug("Attempting OCR for page %d in %s", page_num, file_path.name)
            try:
                images = convert_from_path(str(file_path), first_page=page_num, last_page=page_num)
                if images:
                    ocr_text = pytesseract.image_to_string(images[0]).strip()
                    if ocr_text:
                        text = ocr_text
            except Exception as e:
                logger.warning("OCR failed for page %d in %s: %s", page_num, file_path.name, e)

        # Skip blank pages — they produce empty chunks and waste embeddings
        if not text:
            logger.debug("Skipping blank page %d in %s", page_num, file_path.name)
            continue

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source":      file_path.name,
                    "file_type":   "pdf",
                    "page":        page_num,
                    "total_pages": len(reader.pages),
                }
            )
        )

    return documents


def _load_docx(file_path: Path) -> List[Document]:
    """
    Parse a DOCX file using python-docx.

    A DOCX file is a ZIP archive containing XML files.
    python-docx extracts paragraph text automatically.

    WHY one Document for the whole file:
    DOCX files don't have a reliable page concept in their XML.
    We return the full text and let the chunker split it.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = DocxDocument(str(file_path))

    paragraphs = [
        para.text.strip()
        for para in doc.paragraphs
        if para.text.strip()
    ]

    full_text = "\n\n".join(paragraphs)

    if not full_text:
        logger.warning("DOCX file appears to be empty: %s", file_path.name)
        return []

    return [
        Document(
            page_content=full_text,
            metadata={
                "source":    file_path.name,
                "file_type": "docx",
                "page":      1,
            }
        )
    ]


def _load_txt(file_path: Path) -> List[Document]:
    """
    Parse a plain text file.

    Handles encoding gracefully — UTF-8 with fallback to latin-1.

    WHY two encoding attempts:
    Many corporate documents use Windows-1252 or latin-1 encoding.
    UTF-8 decode raises UnicodeDecodeError on these files.
    latin-1 fallback handles 99% of real-world cases.
    """
    try:
        text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        logger.warning(
            "UTF-8 decode failed for %s, trying latin-1", file_path.name
        )
        text = file_path.read_text(encoding="latin-1")

    text = text.strip()
    if not text:
        logger.warning("TXT file appears to be empty: %s", file_path.name)
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source":    file_path.name,
                "file_type": "txt",
                "page":      1,
            }
        )
    ]


def _load_csv(file_path: Path) -> List[Document]:
    """
    Parse a CSV file using pandas.

    WHY treat each row as a Document:
    A CSV row is a discrete record (one product, one transaction).
    Each row should be independently searchable.

    We convert each row to a human-readable string:
    "column_name: value | column_name: value | ..."

    This format embeds well because it reads like natural language.
    """
    try:
        import pandas as pd
    except ImportError:
        raise ImportError("pandas not installed. Run: pip install pandas")

    df = pd.read_csv(str(file_path))
    df = df.dropna(how="all")

    if df.empty:
        logger.warning("CSV file appears to be empty: %s", file_path.name)
        return []

    documents = []
    for row_idx, row in df.iterrows():
        row_text = " | ".join(
            f"{col}: {val}"
            for col, val in row.items()
            if pd.notna(val)
        )

        documents.append(
            Document(
                page_content=row_text,
                metadata={
                    "source":    file_path.name,
                    "file_type": "csv",
                    "row":       int(row_idx) + 1,
                    "page":      1,
                }
            )
        )

    return documents